import praw
from docx import Document
from docx.shared import Pt
from tqdm import tqdm
from colorama import init, Fore, Style
import time
import re
from openai import OpenAIError
import os
import openai
from dotenv import load_dotenv

load_dotenv()

# Clear the terminal screen
def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

clear_screen()

openai.api_key = os.getenv("open_api_key")

# Initialize the Reddit Client
reddit = praw.Reddit(
    client_id=os.getenv("reddit_client_id"),
    client_secret=os.getenv("reddit_client_secret"),
    user_agent='script by @tylergbbs',
)

print(Fore.LIGHTYELLOW_EX + "========== Reddit Post Fetcher and Rewriter ==========" + Style.RESET_ALL)
print()

# Inputs
num_of_posts = int(input(Fore.LIGHTBLUE_EX + "How many posts do you want to fetch? " + Style.RESET_ALL))
clear_screen()
num_of_variations = int(input(Fore.LIGHTBLUE_EX + "How many variations of each post do you want? " + Style.RESET_ALL))
clear_screen()
subreddit_name = input(Fore.LIGHTBLUE_EX + "Enter a subreddit to search: " + Style.RESET_ALL)
clear_screen()

print(Style.BRIGHT + Fore.LIGHTBLUE_EX + f"Fetching and rewriting {num_of_posts} posts, each with {num_of_variations} variations." + Style.RESET_ALL)
if subreddit_name:
    print(f"Subreddit: {subreddit_name}")

# Create a new directory for saving files
os.makedirs('posts', exist_ok=True)

def fetch_reddit_posts():
    subreddit = reddit.subreddit(subreddit_name)
    posts = subreddit.hot(limit=num_of_posts)

    urls = []  # List to store the post URLs
    for post in posts:
        urls.append(f"https://reddit.com{post.permalink}")

    return urls  # Return the list of post URLs

post_urls = fetch_reddit_posts()

def sanitize_title(title):
    return "".join(c for c in title if c.isalnum() or c.isspace())

def generate_post_content_with_retry(post_url: str):
    retries = 0
    while retries < 10:  # Maximum 10 retries
        try:
            messages = []
            messages.append({"role": "user", "content": f"""I will give you a URL to a Reddit post. Your job is to rewrite the post and make it more engaging in a article format. All I want is the post, no extra comments. POST URL: {post_url}"""})

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            print("API Error:", e)
            time.sleep(1)
            clear_screen()
            print("Retrying in 15 seconds...")
            time.sleep(15)
            retries += 1

def generate_post_title_with_retry(post_content):
    retries = 0
    while retries < 3:  # Maximum 3 retries
        try:
            messages = []
            messages.append({"role": "user", "content": f"I will give you a Reddit post. Your job is to give it an engaging and clickable title. All I want is the title, no extra comments. Here is the post: {post_content}"})

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            print("API Error:", e)
            time.sleep(1)
            clear_screen()
            print("Retrying in 15 seconds...")
            time.sleep(15)
            retries += 1

# Iterate over the posts
for url in tqdm(post_urls, desc="Rewriting Posts", unit="post", colour="green"):
    # Rewrite the post three times
    for i in range(num_of_variations):
        try:
            # Generate rewritten content
            rewritten_content = generate_post_content_with_retry(url)

            # Generate the title
            title = generate_post_title_with_retry(rewritten_content)
            
            # Sanitize the title to be used as filename
            sanitized_title = sanitize_title(title)

            # Initialize the document
            document = Document()
            
            # Add the title as a heading
            title_heading = document.add_heading(level=1)
            title_heading.add_run(title).font.size = Pt(14)
            
            # Add the rewritten content as a paragraph
            content_paragraph = document.add_paragraph()
            content_paragraph.add_run(rewritten_content).font.size = Pt(8)

            # Add an empty paragraph for spacing
            document.add_paragraph()

            # Save each post as a separate document
            document.save(f"posts/{sanitized_title}_variation_{i + 1}.docx")

            # Delay before the next API request
            # time.sleep(10)
        except Exception as e:
            print("Error occurred:", e)

clear_screen()
print(Fore.LIGHTGREEN_EX + "Reddit Post rewriting completed." + Style.RESET_ALL)



#notes
#apply multi-threading to make program faster
#apply modularity
#minimize the functions
#optimize their usage
#tailor prompt to produce the best output


#Brainstorming features to add 
#fetch from multiple subreddits
#add history log (for post that is already written. To avoid rewritting the same post. That's if you plan to run within >2 hours of each other so other post have time to show)